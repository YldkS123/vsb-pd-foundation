# -*- coding: utf-8 -*-
"""Build paper_draft.docx from docs/paper_draft.md.

Design system: narrative_proposal preset (see documents skill design presets)
with a single named override: CJK fonts (SimSun body / SimHei headings) so
Chinese text renders correctly. US Letter, 1in margins, 9360 DXA content width.
Lists use real Word numbering definitions created in the numbering part.
Tables use fixed DXA geometry (tblW/tblInd/tblGrid/tcW agree).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "paper_draft.md"
OUT = ROOT / "docs" / "paper_draft.docx"

# ---------------------------------------------------------------- tokens ----
BODY_LATIN = "Calibri"
BODY_EAST = "SimSun"          # 宋体
HEAD_LATIN = "Calibri"
HEAD_EAST = "SimHei"          # 黑体
CODE_LATIN = "Consolas"
H1_C = RGBColor(0x2E, 0x74, 0xB5)
H2_C = RGBColor(0x2E, 0x74, 0xB5)
H3_C = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FILL = "F4F6F9"
CODE_FILL = "F4F6F9"

CAPTIONS = {
    0: "表 1  物理特征类别（每个窗口 58 维）",
    1: "表 2  相位感知交互与层级损失消融（开发集 5 折，seed 42）",
    2: "表 3  传统机器学习基线对比（开发集 5 折逐折均值 PR-AUC，K=8 特征缓存 + seed 42）",
    3: "表 4  匹配编码器对比（统一 context-concat 协议，seed 42）",
    4: "表 5  已发表方法基线（历史协议，仅文献位置参考）",
    5: "表 6  编码器机制验证（203k 参考架构）",
    6: "表 7  MIL 聚合机制验证（203k 参考架构）",
    7: "表 8  相位交互机制验证（203k 参考架构）",
    8: "表 9  测量不完善鲁棒性扰动（历史锁定主线，开发集 OOF；203k 仅作机制参照）",
    9: "表 10  时间偏移鲁棒性：历史主线 vs 增强变体（开发集 OOF）",
    10: "表 11  窗口策略消融（203k 机制验证 + 历史轻量主线证据）",
    11: "表 12  外部跨域迁移性分析（冻结外部数据集，历史轻量主线）",
    12: "表 13  采样受控实验（开发集 5 折，seed 42；统一 E4 架构，仅采样策略不同）",
    13: "表 14  采样策略配对 cluster bootstrap 差异（mixed K=8 减各策略，95% CI）",
    14: "表 15  端到端测量→判决管线成本（Stage 2 基准，50 次重复 p50）",
    15: "表 16  测量聚类统计不确定性分析（E1-E6 与匹配编码器配对 bootstrap，开发集 OOF，seed 42）",
    16: "表 17  Harvard Dataverse 独立保留集盲测结果（E4 冻结模型，seed 42 5 折平均）",
}
# ---------------------------------------------------------------- helpers ---
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)$")


def set_font(run, east: str, latin: str, size, color=None, bold=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_inline(par, text: str, size, east=BODY_EAST, latin=BODY_LATIN, color=None):
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = par.add_run(part[2:-2])
            set_font(run, east, latin, size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            set_font(run, BODY_EAST, CODE_LATIN, Pt(size.pt - 1), color=color)
        else:
            run = par.add_run(part)
            set_font(run, east, latin, size, color=color)


def para(doc, text="", *, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         before=0.0, after=8.0, line=1.333, keep_next=False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if keep_next:
        pf.keep_with_next = True
    if text:
        add_inline(p, text, Pt(11))
    return p


def add_figure(doc, alt: str, rel: str):
    """Insert a centered picture with a 9pt SimHei bold caption."""
    img = (SRC.parent / rel).resolve()
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    p.add_run().add_picture(str(img), width=Inches(CONTENT_DXA / 1440))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(alt)
    set_font(cr, HEAD_EAST, HEAD_LATIN, Pt(9), color=None, bold=True)


def shading(el_pr, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    el_pr.append(shd)


def add_numbering(doc, fmt="decimal", text="%1.", start=1):
    """Create an independent numbering definition; return numId."""
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    abstract_id = (max(existing_abstract) + 1) if existing_abstract else 0
    num_id = 1
    existing_nums = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    while num_id in existing_nums:
        num_id += 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start"); start_el.set(qn("w:val"), str(start)); lvl.append(start_el)
    fmt_el = OxmlElement("w:numFmt"); fmt_el.set(qn("w:val"), fmt); lvl.append(fmt_el)
    text_el = OxmlElement("w:lvlText"); text_el.set(qn("w:val"), text); lvl.append(text_el)
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")      # 0.375 in
    ind.set(qn("w:hanging"), "280")   # ~0.194 in
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), BODY_LATIN); rf.set(qn("w:hAnsi"), BODY_LATIN); rf.set(qn("w:eastAsia"), BODY_EAST)
    rpr.append(rf)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def set_para_numbering(p, num_id: int):
    ppr = p._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numpr.append(ilvl)
    numid = numpr.find(qn("w:numId"))
    if numid is None:
        numid = OxmlElement("w:numId"); numpr.append(numid)
    numid.set(qn("w:val"), str(num_id))


def list_item(doc, text, num_id, *, size=11.0, after=4.0, line=1.208):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Inches(0.375)
    pf.first_line_indent = Inches(-0.194)
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    set_para_numbering(p, num_id)
    add_inline(p, text, Pt(size))
    return p


def set_cell_margins(table, top=80, bottom=80, left=120, right=120):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for tag, val in (("w:top", top), ("w:bottom", bottom), ("w:left", left), ("w:right", right)):
        el = mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def eff_len(s: str) -> int:
    return sum(2 if ord(ch) > 0x2E80 else 1 for ch in s)


def col_widths(rows, total=CONTENT_DXA):
    n = len(rows[0])
    weights = []
    for c in range(n):
        m = max(eff_len(row[c]) for row in rows)
        weights.append(max(min(m, 46), 5))
    s = sum(weights)
    widths = [max(640, int(total * w / s)) for w in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tblw = tbl_pr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW"); tbl_pr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths))); tblw.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd"); tbl_pr.append(ind)
    ind.set(qn("w:w"), str(TABLE_INDENT_DXA)); ind.set(qn("w:type"), "dxa")
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_margins(table)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid"); tbl.insert(0, grid)
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row, width in zip(table.rows, [widths] * len(table.rows)):
        for cell, w in zip(row.cells, width):
            tc_pr = cell._tc.get_or_add_tcPr()
            tcw = tc_pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW"); tc_pr.append(tcw)
            tcw.set(qn("w:w"), str(w)); tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(caption)
    set_font(run, HEAD_EAST, HEAD_LATIN, Pt(10.5), color=None, bold=True)

    n_rows, n_cols = len(rows), len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = doc.styles["Table Grid"]
    widths = col_widths(rows)
    set_table_geometry(table, widths)

    for r, row in enumerate(rows):
        tr = table.rows[r]._tr
        if r == 0:
            tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.line_spacing = 1.15
            short = eff_len(val) <= 14
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if short else WD_ALIGN_PARAGRAPH.LEFT
            if r == 0:
                cell_pr = cell._tc.get_or_add_tcPr()
                shading(cell_pr, HEADER_FILL)
                add_inline(p, val, Pt(10.5), east=HEAD_EAST, latin=HEAD_LATIN)
                for rn in p.runs:
                    rn.bold = True
            else:
                add_inline(p, val, Pt(10.5))
    return table


# ---------------------------------------------------------------- parsing ---
def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build():
    lines = SRC.read_text(encoding="utf-8").splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    # Normal style ------------------------------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = BODY_LATIN
    normal.font.size = Pt(11)
    rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_LATIN)
    rfonts.set(qn("w:hAnsi"), BODY_LATIN)
    rfonts.set(qn("w:eastAsia"), BODY_EAST)
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(8)
    npf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, H1_C, 18, 10),
        ("Heading 2", 13, H2_C, 12, 6),
        ("Heading 3", 12, H3_C, 8, 4),
    ):
        st = doc.styles[name]
        st.font.name = HEAD_LATIN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        rf = st.element.get_or_add_rPr().get_or_add_rFonts()
        rf.set(qn("w:ascii"), HEAD_LATIN)
        rf.set(qn("w:hAnsi"), HEAD_LATIN)
        rf.set(qn("w:eastAsia"), HEAD_EAST)
        pf = st.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True

    # header / footer ----------------------------------------------------------
    header_p = sec.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header_p.add_run("覆盖感知采样与层级弱监督局部放电检测 · 论文修订稿")
    set_font(hr, HEAD_EAST, HEAD_LATIN, Pt(9), color=MUTED)
    ppr = header_p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom); ppr.append(pbdr)

    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer_p.add_run()
    set_font(fr, BODY_EAST, BODY_LATIN, Pt(9), color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    inner_r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "1"
    inner_r.append(t)
    fld.append(inner_r)
    footer_p._p.append(fld)

    # numbering ----------------------------------------------------------------
    num_bullet = add_numbering(doc, fmt="bullet", text="•")

    # content ------------------------------------------------------------------
    i = 0
    n = len(lines)
    table_idx = 0
    list_stack = []  # ('bullet'|'decimal', num_id)
    prev_num = 0          # last decimal number in the active list
    cur_decimal_id = None # numbering instance of the active decimal list
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("```"):
            # code block until closing fence
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            block = doc.add_paragraph()
            bpf = block.paragraph_format
            bpf.left_indent = Inches(0.15)
            bpf.space_before = Pt(4)
            bpf.space_after = Pt(8)
            bpf.line_spacing = 1.0
            shading(block._p.get_or_add_pPr(), CODE_FILL)
            for k, cl in enumerate(code_lines):
                if k > 0:
                    block.add_run().add_break()
                r = block.add_run(cl if cl else " ")
                set_font(r, BODY_EAST, CODE_LATIN, Pt(9.5))
            prev_num = 0
            cur_decimal_id = None
            continue
        if stripped.startswith(">"):
            body = stripped.lstrip(">").strip()
            p = para(doc, body, align=WD_ALIGN_PARAGRAPH.LEFT, after=3, line=1.15)
            for r in p.runs:
                r.font.color.rgb = MUTED
                r.font.size = Pt(10)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        if stripped.startswith("# "):
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(14)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(title_text)
            set_font(r, HEAD_EAST, HEAD_LATIN, Pt(18), color=None, bold=True)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:].strip(), Pt(16), east=HEAD_EAST, latin=HEAD_LATIN)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:].strip(), Pt(13), east=HEAD_EAST, latin=HEAD_LATIN)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:].strip(), Pt(12), east=HEAD_EAST, latin=HEAD_LATIN)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows, CAPTIONS[table_idx])
                table_idx += 1
            prev_num = 0
            cur_decimal_id = None
            continue
        if re.match(r"^\d+\.\s", stripped):
            num = int(stripped.split(".")[0])
            text = re.sub(r"^\d+\.\s*", "", stripped)
            if cur_decimal_id is None or num != prev_num + 1:
                cur_decimal_id = add_numbering(doc, fmt="decimal", text="%1.")
            prev_num = num
            num_id = cur_decimal_id
            list_item(doc, text, num_id)
            i += 1
            continue
        if stripped.startswith("- "):
            if not list_stack or list_stack[-1][0] != "bullet":
                list_stack.append(("bullet", num_bullet))
            list_item(doc, stripped[2:].strip(), num_bullet)
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
    # plain paragraph
        m = IMG_RE.match(stripped)
        if m:
            add_figure(doc, m.group("alt"), m.group("src"))
            prev_num = 0
            cur_decimal_id = None
            i += 1
            continue
        para(doc, stripped)
        prev_num = 0
        cur_decimal_id = None
        i += 1


    props = doc.core_properties
    props.title = "面向超长稀疏三相局部放电的覆盖感知采样与轻量层级弱监督检测"
    props.author = "（作者待填）"
    props.subject = "VSB 局部放电检测 · 论文修订稿"

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
