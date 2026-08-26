# -*- coding: utf-8 -*-
"""Build paper_tii.docx from docs/paper_tii_v2.md.

English TII manuscript DOCX: Calibri body/headings, US Letter, 1in margins,
IEEE TII double-column compatible. Tables, headings, paragraphs handled.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "paper_tii_v2.md"
OUT = ROOT / "docs" / "paper_tii.docx"

BODY_LATIN = "Calibri"
HEAD_LATIN = "Calibri"
CODE_LATIN = "Consolas"
H1_C = RGBColor(0x2E, 0x74, 0xB5)
H2_C = RGBColor(0x2E, 0x74, 0xB5)
H3_C = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FILL = "F4F6F9"
CODE_FILL = "F4F6F9"


def set_font(run, latin=BODY_LATIN, east=None, size=10.5, bold=False, italic=False, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    if east:
        rFonts.set(qn("w:eastAsia"), east)


def add_par(doc, text, size=10.5, bold=False, italic=False, align=None,
            space_before=0, space_after=6, color=None, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if keep:
        pf.keep_with_next = True
    # handle **bold** and *italic* inline
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, size=size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_font(r, size=size, italic=True, color=color)
        else:
            r = p.add_run(part)
            set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    color = H1_C if level == 1 else (H2_C if level == 2 else H3_C)
    size = 14 if level == 1 else (12 if level == 2 else 11)
    r = p.add_run(text)
    set_font(r, latin=HEAD_LATIN, size=size, bold=True, color=color)
    return p


def add_table(doc, rows, caption=None):
    if caption:
        add_par(doc, caption, size=9, italic=True, space_before=8, space_after=2)
    ncol = max(len(r) for r in rows)
    nrow = len(rows)
    table = doc.add_table(rows=nrow, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            txt = row[j] if j < len(row) else ""
            r = p.add_run(txt)
            set_font(r, size=9, bold=(i == 0))
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), HEADER_FILL)
                tcPr.append(shd)
    add_par(doc, "", size=4, space_after=2)
    return table


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    # page setup: US Letter, 1in margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(1))

    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_rows = []
    table_caption = None
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        # skip blockquote header (author/status notes)
        if s.startswith(">") and not s.startswith(">"):
            i += 1
            continue
        if s.startswith("> "):
            add_par(doc, s[2:], size=9, italic=True, color=MUTED, space_after=2)
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        # table caption (bold TABLE ...)
        mcap = re.match(r"\*\*TABLE ([IVX]+(?:-[A-Z])?)\*\*\s*(.*)", s)
        if mcap:
            table_caption = f"TABLE {mcap.group(1)}  {mcap.group(2)}"
            i += 1
            continue
        # table rows
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        if table_rows:
            add_table(doc, table_rows, table_caption)
            table_rows = []
            table_caption = None
            i += 1
            continue
        # code fence
        if s.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                set_font(r, latin=CODE_LATIN, size=9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        # headings
        if s.startswith("### "):
            add_heading(doc, s[4:], 3)
            i += 1
            continue
        if s.startswith("## "):
            add_heading(doc, s[3:], 2)
            i += 1
            continue
        if s.startswith("# "):
            add_heading(doc, s[2:], 1)
            i += 1
            continue
        # list items
        if re.match(r"^\d+\.\s+", s):
            txt = re.sub(r"^\d+\.\s+", "", s)
            p = add_par(doc, txt, size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        if s.startswith("- "):
            p = add_par(doc, s[2:], size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        # image placeholder
        if s.startswith("![") and "](" in s:
            add_par(doc, s, size=9, italic=True, color=MUTED)
            i += 1
            continue
        # blank
        if not s:
            i += 1
            continue
        # equation ($$...$$ on one line or spanning lines)
        if s.startswith("$$"):
            if s.endswith("$$") and len(s) > 4:
                # single-line equation: strip the $$ markers
                eq = s[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(eq)
                set_font(r, latin="Cambria Math", size=10.5, italic=True)
                i += 1
                continue
            # multi-line equation
            i += 1
            eq_buf = []
            while i < len(lines):
                inner = lines[i].strip()
                if inner.endswith("$$"):
                    eq_buf.append(inner[:-2].strip())
                    i += 1
                    break
                eq_buf.append(inner)
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(" ".join(eq_buf))
            set_font(r, latin="Cambria Math", size=10.5, italic=True)
            continue
        # regular paragraph
        add_par(doc, s, size=10.5)
        i += 1

    # flush trailing table
    if table_rows:
        add_table(doc, table_rows, table_caption)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
