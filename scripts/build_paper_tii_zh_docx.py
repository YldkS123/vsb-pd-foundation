# -*- coding: utf-8 -*-
"""Build paper_tii_zh.docx from docs/paper_tii_zh.md with Chinese fonts
and TII figures inserted at matching sections."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "paper_tii_zh.md"
OUT = ROOT / "docs" / "paper_tii_zh.docx"
FIGDIR = ROOT / "docs" / "figures_tii"

BODY_EAST = "宋体"
HEAD_EAST = "黑体"
BODY_LATIN = "Times New Roman"
CODE_LATIN = "Consolas"
H1_C = RGBColor(0x2E, 0x74, 0xB5)
H2_C = RGBColor(0x2E, 0x74, 0xB5)
H3_C = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F4F6F9"

FIG_MAP = [
    ("## 三、所提框架", [
        ("fig1_framework.png", "图 1 框架总览：三相感知 → CAS 采样 → TFE/LPT 编码 → 注意力 MIL → context-concat → noisy-OR，标注三成本轴"),
    ]),
    ("### C. 感知成本优化（采样率）", [
        ("fig2_sampling_rate.png", "图 2 采样率成本-性能曲线（5 MHz 保持 83%）"),
    ]),
    ("### E. 外部多数据集验证（figshare 24033225）", [
        ("fig3_encoder_scale.png", "图 3 编码器对比：VSB vs 外部数据集——数据规模假设"),
    ]),
    ("### D. 标注成本量化", [
        ("fig4_labeling_cost.png", "图 4 标注成本曲线（叠加 VICReg 负结果）"),
    ]),
]


def set_cjk_font(run, east=BODY_EAST, latin=BODY_LATIN, size=10.5, bold=False,
                 italic=False, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), east)


def add_par(doc, text, size=10.5, bold=False, italic=False, align=None,
            space_before=0, space_after=6, color=None, east=BODY_EAST):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_cjk_font(r, east=east, size=size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_cjk_font(r, east=east, size=size, italic=True, color=color)
        else:
            r = p.add_run(part)
            set_cjk_font(r, east=east, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 1 else 12)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    color = H1_C if level == 1 else (H2_C if level == 2 else H3_C)
    size = 16 if level == 1 else (14 if level == 2 else 12)
    r = p.add_run(text)
    set_cjk_font(r, east=HEAD_EAST, size=size, bold=True, color=color)
    return p


def add_table(doc, rows, caption=None):
    if caption:
        add_par(doc, caption, size=9, italic=True, space_before=8, space_after=2)
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            txt = row[j] if j < len(row) else ""
            r = p.add_run(txt)
            set_cjk_font(r, size=9, bold=(i == 0))
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), HEADER_FILL)
                tcPr.append(shd)
    add_par(doc, "", size=4, space_after=2)
    return table


def add_figure(doc, img_name, caption):
    img_path = FIGDIR / img_name
    if not img_path.exists():
        add_par(doc, f"[图缺失: {img_name}]", size=9, italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_cjk_font(r, size=9, bold=False, italic=True, color=MUTED)


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(1))

    fig_lookup = {h: figs for h, figs in FIG_MAP}

    lines = md.splitlines()
    i = 0
    table_rows = []
    table_caption = None
    inserted = set()

    def flush_table():
        nonlocal table_rows, table_caption
        if table_rows:
            add_table(doc, table_rows, table_caption)
            table_rows = []
            table_caption = None

    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if s.startswith("> "):
            add_par(doc, s[2:], size=9, italic=True, color=MUTED, space_after=2)
            i += 1
            continue
        if s == "---":
            flush_table()
            i += 1
            continue
        mcap = re.match(r"\*\*表 ([IVX]+(?:-[A-Z])?)\*\*\s*(.*)", s)
        if mcap:
            table_caption = f"表 {mcap.group(1)}  {mcap.group(2)}"
            i += 1
            continue
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        flush_table()
        if s.startswith("### "):
            add_heading(doc, s[4:], 3)
            for h, figs in fig_lookup.items():
                if h in s and h not in inserted:
                    for img, cap in figs:
                        add_figure(doc, img, cap)
                    inserted.add(h)
            i += 1
            continue
        if s.startswith("## "):
            add_heading(doc, s[3:], 2)
            for h, figs in fig_lookup.items():
                if h in s and h not in inserted:
                    for img, cap in figs:
                        add_figure(doc, img, cap)
                    inserted.add(h)
            i += 1
            continue
        if s.startswith("# "):
            add_heading(doc, s[2:], 1)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", s):
            txt = re.sub(r"^\d+\.\s+", "", s)
            p = add_par(doc, txt, size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        if s.startswith("- "):
            p = add_par(doc, s[2:], size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        if s.startswith("!["):
            i += 1
            continue
        if not s:
            i += 1
            continue
        if s.startswith("$$"):
            if s.endswith("$$") and len(s) > 4:
                eq = s[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(eq)
                set_cjk_font(r, east="宋体", latin="Cambria Math", size=10.5, italic=True)
                i += 1
                continue
            i += 1
            eq_buf = []
            while i < len(lines):
                inner = lines[i].strip()
                if inner.endswith("$$"):
                    eq_buf.append(inner[:-2].strip())
                    i += 1
                    break
                eq_buf.append(inner)
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(" ".join(eq_buf))
            set_cjk_font(r, east="宋体", latin="Cambria Math", size=10.5, italic=True)
            continue
        add_par(doc, s, size=10.5)
        i += 1

    flush_table()
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Figures inserted: {len(inserted)} sections")


if __name__ == "__main__":
    main()
