# -*- coding: utf-8 -*-
"""Build IEEE TIM double-column .docx from the English markdown paper.

Produces a two-column US Letter document with IEEE-style formatting:
- 0.75in margins (IEEE standard for TRANS journals)
- Two-column layout starting from the body
- Title, author block, abstract single-column spanning both columns
- 10pt body text (IEEE TRANS standard)
- Numbered sections, centered captions, figures
- Page numbers in footer
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor, Cm, Emu

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "paper_draft_en.md"
OUT = ROOT / "docs" / "paper_draft_ieee.docx"

# ---------------------------------------------------------------- tokens ----
BODY_LATIN = "Times New Roman"
HEAD_LATIN = "Times New Roman"
CODE_LATIN = "Courier New"
TITLE_FONT = "Times New Roman"
H1_C = RGBColor(0, 0, 0)
H2_C = RGBColor(0, 0, 0)
H3_C = RGBColor(0, 0, 0)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F0F0F0"
CODE_FILL = "F5F5F5"

IEEE_MARGIN = Inches(0.75)       # IEEE standard
COLUMN_GAP = Inches(0.25)        # gap between columns

CAPTIONS_EN = {
    0: "Table 1  Physical feature categories (58 dimensions per window)",
    1: "Table 2  Traditional machine learning baselines (development set 5-fold per-fold mean PR-AUC, K=8 feature cache + seed 42)",
    2: "Table 3  Phase interaction and hierarchical loss ablation (E1\u2013E6, development set 5-fold, seed 42)",
    3: "Table 4  Matched encoder comparison (unified context-concat protocol, seed 42)",
    4: "Table 5  Published method baselines (historical protocol, literature reference only)",
    5: "Table 6  Encoder mechanism verification (203k reference architecture)",
    6: "Table 7  MIL aggregation mechanism verification (203k reference architecture)",
    7: "Table 8  Phase interaction mechanism verification (203k reference architecture)",
    8: "Table 9  Measurement imperfection robustness perturbation (historical locked mainline, development set OOF)",
    9: "Table 10  Time shift robustness: historical mainline vs augmented variant (development set OOF)",
    10: "Table 11  Window strategy ablation (203k mechanism verification + historical lightweight mainline evidence)",
    11: "Table 12  External cross-domain transferability analysis (frozen external datasets, historical lightweight mainline)",
    12: "Table 13  Controlled sampling experiment (development set 5-fold, seed 42; unified E4 architecture, sampling strategy only)",
    13: "Table 14  Sampling contrast bootstrap: paired measurement-clustered differences (mixed K=8 minus each alternative, development set OOF, seed 42)",
    14: "Table 15  End-to-end measurement-to-decision pipeline cost (Stage 2 benchmark, 50 repetitions p50)",
    15: "Table 16  Harvard Dataverse independent held-out set blind test results (E4 frozen model, seed 42 5-fold average)",
    16: "Table 17  Measurement-cluster statistical uncertainty analysis (E1\u2013E6 and matched encoder paired bootstrap, development set OOF, seed 42)",
}

# ---------------------------------------------------------------- helpers ---
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)$")


def set_font(run, latin: str, size, color=None, bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), latin)
    run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(par, text: str, size, latin=BODY_LATIN, color=None):
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = par.add_run(part[2:-2])
            set_font(run, latin, size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            set_font(run, CODE_LATIN, Pt(size.pt - 1), color=color)
        else:
            run = par.add_run(part)
            set_font(run, latin, size, color=color)


def para(doc, text="", *, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         before=0.0, after=4.0, line=1.05, keep_next=False, size=Pt(10)):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if keep_next:
        pf.keep_with_next = True
    if text:
        add_inline(p, text, size)
    return p


def shading(el_pr, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    el_pr.append(shd)


def set_cell_font(cell, text: str, size=Pt(8.5), bold=False, latin=BODY_LATIN):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, latin, size, bold=bold)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge, val in kwargs.items():
        tag = qn(f"w:{edge}")
        el = tcPr.find(tag)
        if el is None:
            el = OxmlElement(tag)
            tcPr.append(el)
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), val.get("color", "000000"))


def add_numbering(doc, fmt="decimal", text="%1."):
    """Create a numbering definition and return its numId."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    num_id = 1
    existing_ids = []
    for child in numbering:
        if child.tag == qn("w:num"):
            existing_ids.append(int(child.get(qn("w:numId"))))
    num_id = max(existing_ids) + 1 if existing_ids else 1
    ab_el = OxmlElement("w:abstractNum")
    ab_el.set(qn("w:abstractNumId"), str(num_id + 1000))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), fmt)
    lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), text)
    lvl.append(lvlText)
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "ListParagraph")
    lvl.append(pStyle)
    ab_el.append(lvl)
    numbering.append(ab_el)
    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(num_id))
    abSrc = OxmlElement("w:abstractNumId")
    abSrc.set(qn("w:val"), str(num_id + 1000))
    num_el.append(abSrc)
    numbering.append(num_el)
    return num_id


def parse_table(lines, start):
    """Return (list_of_lists_of_cells, next_index)."""
    rows = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith("|"):
            break
        cells = [c.strip() for c in stripped.split("|")]
        # Remove first/last empty from leading/trailing |
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        # Skip separator rows
        if all(c == "---" or "---" in c for c in cells if c):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows, caption):
    """Add a table with caption."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            text = row_data[ci] if ci < len(row_data) else ""
            is_header = ri == 0
            set_cell_font(cell, text, size=Pt(8), bold=is_header)
            if is_header:
                shading(cell._tc.get_or_add_tcPr(), HEADER_FILL)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set column widths proportionally
    col_width = int(Inches(6.5 / ncols))
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    # Table caption
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = True
    cr = cap.add_run(caption)
    set_font(cr, HEAD_LATIN, Pt(9), bold=True)

    # Space after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(2)
    sp.paragraph_format.line_spacing = 0.2


def add_figure(doc, alt: str, rel: str):
    """Insert a centered picture with caption."""
    img = (SRC.parent / rel).resolve()
    if not img.exists():
        para(doc, f"[Figure not found: {img.name}]", size=Pt(10))
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    # Scale to column width (3.25in for two-column)
    p.add_run().add_picture(str(img), width=Inches(3.0))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(6)
    cr = cap.add_run(alt)
    set_font(cr, HEAD_LATIN, Pt(9), bold=True)


def set_section_properties(section, num_cols=1):
    """Configure section page layout and column settings."""
    sect_pr = section._sectPr

    # Page size: US Letter
    pgSz = sect_pr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sect_pr.append(pgSz)
    pgSz.set(qn("w:w"), "12240")   # 8.5in in twips
    pgSz.set(qn("w:h"), "15840")   # 11in in twips

    # Margins
    for margin, val in [("top", "1440"), ("bottom", "1440"),
                        ("left", "1080"), ("right", "1080"),
                        ("header", "720"), ("footer", "720")]:
        el = sect_pr.find(qn(f"w:{margin}"))
        if el is None:
            el = OxmlElement(f"w:{margin}")
            sect_pr.append(el)
        el.set(qn("w:w"), val)

    if num_cols > 1:
        # Remove existing cols element if any
        for old in sect_pr.findall(qn("w:cols")):
            sect_pr.remove(old)
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), str(num_cols))
        cols.set(qn("w:space"), "360")  # 0.25in gap in twips
        # Equal-width columns
        cols.set(qn("w:equalWidth"), "true")
        # Insert at beginning of sectPr
        sect_pr.insert(0, cols)


def setup_document():
    """Create document with proper section setup."""
    doc = Document()

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # First section properties
    section = doc.sections[0]
    set_section_properties(section, num_cols=1)

    # Default paragraph format
    style = doc.styles['Normal']
    style.font.name = BODY_LATIN
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.05

    return doc


def make_columns(doc, section, num_cols=2):
    """Make a section multi-column by modifying sectPr."""
    set_section_properties(section, num_cols=num_cols)


def build():
    with open(SRC, "r", encoding="utf-8") as f:
        text = f.read()
    lines = text.split("\n")

    doc = setup_document()
    section = doc.sections[0]

    # ---------------------------------------------------------------- parse ----
    i = 0
    n = len(lines)
    table_idx = 0
    in_title = True
    in_abstract = False
    body_started = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("```"):
            # Code block
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if body_started:
                block = doc.add_paragraph()
                bpf = block.paragraph_format
                bpf.left_indent = Inches(0.15)
                bpf.space_before = Pt(2)
                bpf.space_after = Pt(4)
                bpf.line_spacing = 0.95
                shading(block._p.get_or_add_pPr(), CODE_FILL)
                for k, cl in enumerate(code_lines):
                    if k > 0:
                        block.add_run().add_break()
                    r = block.add_run(cl if cl else " ")
                    set_font(r, CODE_LATIN, Pt(8))
            continue

        if stripped.startswith(">"):
            if body_started:
                body = stripped.lstrip(">").strip()
                p = para(doc, body, align=WD_ALIGN_PARAGRAPH.LEFT,
                         before=2, after=2, line=1.0, size=Pt(9.5))
                for r in p.runs:
                    r.font.color.rgb = MUTED
            i += 1
            continue

        # ---- Title ----
        if in_title and stripped.startswith("# "):
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(title_text)
            set_font(r, TITLE_FONT, Pt(18), bold=True)
            in_title = False
            in_abstract = True
            i += 1
            continue

        # ---- Section headers ----
        if stripped.startswith("## ") and not in_abstract:
            if not body_started:
                # Start two-column layout after abstract
                body_started = True
                make_columns(doc, section, num_cols=2)
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, stripped[3:].strip(), Pt(11), latin=HEAD_LATIN)
            # Make it bold
            for r in p.runs:
                r.bold = True
            i += 1
            continue

        if stripped.startswith("### "):
            if not body_started:
                body_started = True
                make_columns(doc, section, num_cols=2)
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, stripped[4:].strip(), Pt(10.5), latin=HEAD_LATIN)
            for r in p.runs:
                r.bold = True
            i += 1
            continue

        if stripped.startswith("#### "):
            if not body_started:
                body_started = True
                make_columns(doc, section, num_cols=2)
            p = doc.add_paragraph(style="Heading 3")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, stripped[5:].strip(), Pt(10), latin=HEAD_LATIN)
            for r in p.runs:
                r.bold = True
                r.italic = True
            i += 1
            continue

        # ---- Tables ----
        if stripped.startswith("|") and "|" in stripped[1:]:
            rows, i = parse_table(lines, i)
            if rows:
                if not body_started:
                    body_started = True
                    make_columns(doc, section, num_cols=2)
                add_table(doc, rows, CAPTIONS_EN.get(table_idx, f"Table {table_idx + 1}"))
                table_idx += 1
            continue

        # ---- Figures ----
        m = IMG_RE.match(stripped)
        if m:
            if not body_started:
                body_started = True
                make_columns(doc, section, num_cols=2)
            add_figure(doc, m.group("alt"), m.group("src"))
            i += 1
            continue

        # ---- Ordered lists ----
        om = re.match(r"^\d+\.\s", stripped)
        if om:
            if not body_started:
                body_started = True
                make_columns(doc, section, num_cols=2)
            text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            # Number prefix
            num_part = om.group().strip()
            r = p.add_run(f"{num_part} ")
            set_font(r, BODY_LATIN, Pt(10))
            # Content
            add_inline(p, text, Pt(10))
            i += 1
            continue

        # ---- Bullet lists ----
        if stripped.startswith("- ") or stripped.startswith("* "):
            if not body_started:
                body_started = True
                make_columns(doc, section, num_cols=2)
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run("\u2022 ")
            set_font(r, BODY_LATIN, Pt(10))
            add_inline(p, text, Pt(10))
            i += 1
            continue

        # ---- Abstract ----
        if in_abstract and stripped.startswith("**Abstract"):
            # In the English markdown, abstract may be bold
            text = stripped.replace("**", "")
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_font(r, BODY_LATIN, Pt(10), italic=True)
            in_abstract = False
            i += 1
            continue

        # ---- Regular paragraph ----
        if body_started:
            p = para(doc, stripped, size=Pt(10))
        else:
            # Before body start (abstract-like text)
            p = para(doc, stripped, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     before=2, after=4, line=1.05, size=Pt(10))

        i += 1

    # ---------------------------------------------------------------- footer ----
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)

        # Page number field
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run = fp.add_run()
        run._element.append(fld_char_begin)
        r2 = fp.add_run()
        r2._element.append(instr_text)
        r3 = fp.add_run()
        r3._element.append(fld_char_separate)
        r4 = fp.add_run("1")
        r4.font.name = BODY_LATIN
        r5 = fp.add_run()
        r5._element.append(fld_char_end)

    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    build()
