# -*- coding: utf-8 -*-
"""Build paper_tim_zh.docx from docs/paper_tim_zh_formal.md with Chinese
fonts (SimSun/SimHei) and paper figures inserted at the matching sections.

Figures are inserted after specific section headings:
  三、覆盖感知测量采样  -> fig1_architecture.png (框架图), fig2_window_sampling.png
  四、层级弱监督检测    -> fig1_architecture.png (已插入则不重复)
  五、实验 - E 窗口消融 -> fig8_window_policy.png, fig6_ablation.png
  五、实验 - F 鲁棒性  -> fig7_robustness.png
  五、实验 - G 受控采样 -> fig9_sampling_policy.png
  五、实验 - C 交互消融 -> fig3_pr_curves.png, fig4_reliability.png, fig5_examples.png
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "paper_tim_zh_formal.md"
OUT = ROOT / "docs" / "paper_tim_zh.docx"
FIGDIR = ROOT / "docs" / "figures"

BODY_EAST = "宋体"
HEAD_EAST = "黑体"
BODY_LATIN = "Times New Roman"
CODE_LATIN = "Consolas"
H1_C = RGBColor(0x2E, 0x74, 0xB5)
H2_C = RGBColor(0x2E, 0x74, 0xB5)
H3_C = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F4F6F9"

# figure insertion map: section heading substring -> list of (image, caption)
FIG_MAP = [
    ("### C. 传统机器学习基线与交互消融", [
        ("fig3_pr_curves.png", "图 3 开发集 PR 曲线（历史锁定主模型作对照）"),
        ("fig4_reliability.png", "图 4 盲测可靠性图（历史锁定主模型）"),
        ("fig5_examples.png", "图 5 盲测正/负例示例与注意力权重（历史锁定主模型）"),
    ]),
    ("### E. 窗口策略消融", [
        ("fig8_window_policy.png", "图 8 窗口策略消融（K 与组成）"),
        ("fig6_ablation.png", "图 6 消融实验（历史参考架构机制验证）"),
    ]),
    ("### F. 测量不完善鲁棒性", [
        ("fig7_robustness.png", "图 7 噪声与扰动鲁棒性（历史锁定主模型）"),
    ]),
    ("### G. 受控采样实验：信息选择作为设计变量", [
        ("fig9_sampling_policy.png", "图 9 受控采样实验：策略性能与成本对比"),
    ]),
    ("## 三、覆盖感知测量采样", [
        ("fig1_architecture.png", "图 1 覆盖感知多窗口采样与层级弱监督检测框架"),
        ("fig2_window_sampling.png", "图 2 混合窗口采样示例（测量 705，C 相，窗口标注）"),
    ]),
]


def set_cjk_font(run, east=BODY_EAST, latin=BODY_LATIN, size=10.5, bold=False,
                 italic=False, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), east)


def add_par(doc, text, size=10.5, bold=False, italic=False, align=None,
            space_before=0, space_after=6, color=None, east=BODY_EAST):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_cjk_font(r, east=east, size=size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_cjk_font(r, east=east, size=size, italic=True, color=color)
        else:
            r = p.add_run(part)
            set_cjk_font(r, east=east, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 1 else 12)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    color = H1_C if level == 1 else (H2_C if level == 2 else H3_C)
    size = 16 if level == 1 else (14 if level == 2 else 12)
    r = p.add_run(text)
    set_cjk_font(r, east=HEAD_EAST, size=size, bold=True, color=color)
    return p


def add_table(doc, rows, caption=None):
    if caption:
        add_par(doc, caption, size=9, italic=True, space_before=8, space_after=2)
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            txt = row[j] if j < len(row) else ""
            r = p.add_run(txt)
            set_cjk_font(r, size=9, bold=(i == 0))
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), HEADER_FILL)
                tcPr.append(shd)
    add_par(doc, "", size=4, space_after=2)
    return table


def add_figure(doc, img_name, caption):
    img_path = FIGDIR / img_name
    if not img_path.exists():
        add_par(doc, f"[图缺失: {img_name}]", size=9, italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_cjk_font(r, size=9, bold=False, italic=True, color=MUTED)


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(1))

    # build the figure map lookup by heading substring
    fig_lookup = {h: figs for h, figs in FIG_MAP}

    lines = md.splitlines()
    i = 0
    table_rows = []
    table_caption = None
    inserted_headings = set()

    def flush_table():
        nonlocal table_rows, table_caption
        if table_rows:
            add_table(doc, table_rows, table_caption)
            table_rows = []
            table_caption = None

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if s.startswith("> "):
            add_par(doc, s[2:], size=9, italic=True, color=MUTED, space_after=2)
            i += 1
            continue
        if s == "---":
            flush_table()
            i += 1
            continue
        mcap = re.match(r"\*\*表 ([IVX]+(?:-[A-Z])?)\*\*\s*(.*)", s)
        if mcap:
            table_caption = f"表 {mcap.group(1)}  {mcap.group(2)}"
            i += 1
            continue
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        flush_table()
        if s.startswith("### "):
            add_heading(doc, s[4:], 3)
            for h, figs in fig_lookup.items():
                if h in s and h not in inserted_headings:
                    for img, cap in figs:
                        add_figure(doc, img, cap)
                    inserted_headings.add(h)
            i += 1
            continue
        if s.startswith("## "):
            add_heading(doc, s[3:], 2)
            for h, figs in fig_lookup.items():
                if h in s and h not in inserted_headings:
                    for img, cap in figs:
                        add_figure(doc, img, cap)
                    inserted_headings.add(h)
            i += 1
            continue
        if s.startswith("# "):
            add_heading(doc, s[2:], 1)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", s):
            txt = re.sub(r"^\d+\.\s+", "", s)
            p = add_par(doc, txt, size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        if s.startswith("- "):
            p = add_par(doc, s[2:], size=10.5, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        if s.startswith("!["):
            i += 1
            continue
        if not s:
            i += 1
            continue
        if s.startswith("$$"):
            if s.endswith("$$") and len(s) > 4:
                eq = s[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(eq)
                set_cjk_font(r, east="宋体", latin="Cambria Math", size=10.5, italic=True)
                i += 1
                continue
            i += 1
            eq_buf = []
            while i < len(lines):
                inner = lines[i].strip()
                if inner.endswith("$$"):
                    eq_buf.append(inner[:-2].strip())
                    i += 1
                    break
                eq_buf.append(inner)
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(" ".join(eq_buf))
            set_cjk_font(r, east="宋体", latin="Cambria Math", size=10.5, italic=True)
            continue
        # regular paragraph
        add_par(doc, s, size=10.5)
        i += 1

    flush_table()
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Figures inserted: {len(inserted_headings)} sections")


if __name__ == "__main__":
    main()
