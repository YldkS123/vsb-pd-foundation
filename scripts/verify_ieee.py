"""Verify the IEEE double-column docx."""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'C:\Users\hrfxgfx\Desktop\1112\docs\paper_draft_ieee.docx')

sec = doc.sections[0]
sect_pr = sec._sectPr

print("=== Section Properties ===")
for child in sect_pr:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'cols':
        print(f"  cols: num={child.get(qn('w:num'))}, space={child.get(qn('w:space'))}, equalWidth={child.get(qn('w:equalWidth'))}")
    elif tag == 'pgSz':
        print(f"  page: w={child.get(qn('w:w'))} h={child.get(qn('w:h'))}")
    elif 'mar' in tag:
        print(f"  {tag}={child.get(qn('w:w'))}")

print(f"\n=== Document Stats ===")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")

print(f"\n=== First 10 Paragraphs ===")
for i, p in enumerate(doc.paragraphs[:10]):
    txt = p.text[:100] if p.text else '(empty)'
    print(f"  [{i:3d}] {txt}")

print(f"\n=== Last 5 Paragraphs ===")
for p in doc.paragraphs[-5:]:
    txt = p.text[:120] if p.text else '(empty)'
    print(f"  {txt}")

print(f"\n=== Table Captions (checking all 17) ===")
tbl_count = 0
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt.startswith("Table ") and not txt.startswith("Table "):
        pass
    if txt and txt[0].isdigit() and "Table" in txt[:20]:
        pass
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt.startswith("Table "):
        tbl_count += 1
        print(f"  [{i:3d}] {txt[:90]}")

print(f"\n  Found {tbl_count} table captions")

# Check footer
footer = sec.footer
print(f"\n=== Footer ===")
for p in footer.paragraphs:
    print(f"  text: '{p.text}'")
    
print("\nDONE")
